import pandas as p
import numpy as n
from importlib.metadata import files
from pydoc import doc
from tpyes import authkeys, authpatt,indkeys,prodyear,prodauth,prodauth,prodpubl
import matplotlib.pyplot as pt
from tkinter import *
from tkinter import filedialog
from turtle import bgcolor 
from docx import Document
from docx.shared import Inches
from plots import pbar,pplolt,ppie,pbarh

#ffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff
d = Document()
def todoc(df,picturename,docname):    
    d.add_picture(f'{picturename}.png')
    d.add_page_break()
    l = []
    l.append(list(df.columns))
    for i in df.index:
        l.append(list(df.iloc[i]))     
    table = d.add_table( 0, (len(l[0])))
    for i,j in l:
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(j)
    d.add_page_break()
    d.save(docname)
#fffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff
def prodauth(pa):
     f = p.read_csv(pa)
     t = f['Authors']
     ld = [] # duplicate list of all keywords 
     for i in range(len(t)):     
        temp = str(t.loc[i]).split('; ')
        for j in temp: 
            ki = str(t.loc[i]).split('.,')
            for j in ki:
                ld.append((j.lower()))        
     lu = sorted(list(set(ld)))
     fin = []
     for i in range(len(lu)):
        fin.append([])
        fin[i].append(lu[i])
        fin[i].append(ld.count(lu[i]))
     df = p.DataFrame(fin,columns = ['Author', 'Number of publications'] )
     df.sort_values(by = 'Number of publications' ,inplace = True , ascending = False)
     df = df.head(11) 
     c = list(df.columns)
     df.reset_index(drop = True, inplace = True)     
     return (df)
#------------------------------------------------------------------------------------------------------------------------------
def prodyear(pa):
    f = p.read_csv(pa)
    yearsd = []
    for i in f.index:
        yearsd.append(f.iloc[i,3])
    years = list(n.unique(yearsd))
    tdata = []
    for i in range(len(years)):
        tdata.append([])
        tdata[i].append(str(years[i]))
        tdata[i].append(yearsd.count(years[i]))
    df = p.DataFrame(tdata,columns = ['Year of Publication','Number of Publication'] )                
    c = list(df.columns)
    return (df)
#------------------------------------------------------------------------------------------------------------------------------
def authkeys(pa):
     f = p.read_csv(pa)
     t = f['Author Keywords']
     ld = [] # duplicate list of all keywords 
     for i in range(len(t)):     
        temp = str(t.loc[i]).split('; ')
        for j in temp:
            ld.append((j.lower()))        
     lu = sorted(list(set(ld)))
     lu.remove('nan') 
     fin = []
     for i in range(len(lu)):
        fin.append([])
        fin[i].append(lu[i])
        fin[i].append(ld.count(lu[i]))
     df = p.DataFrame(fin,columns = ['Keywords', 'Occurence'] )
     df.sort_values(by = 'Occurence' ,inplace = True, ascending = False )
     df = df.head(11) 
     c = list(df.columns)
     df.reset_index(drop = True, inplace = True)     
     return (df)
#----------------------------------------------------------------------------------------------------

def indkeys(pa):
     f = p.read_csv(pa)
     t = f['Index Keywords']
     ld = [] 
     for i in range(len(t)):     
        temp = str(t.loc[i]).split('; ')
        for j in temp:
            ld.append((j.lower()))        
     lu = sorted(list(set(ld)))
     lu.remove('nan') 
     fin = []
     for i in range(len(lu)):
        fin.append([])
        fin[i].append(lu[i])
        fin[i].append(ld.count(lu[i]))
     df = p.DataFrame(fin,columns = ['Keywords', 'Occurence'] )
     df.sort_values(by = 'Occurence' ,inplace = True, ascending = False)
     l = []
     c = list(df.columns)
     df = df.head(11) 
     c = list(df.columns)
     df.reset_index(drop = True, inplace = True)
     return (df)
#----------------------------------------------------------------------------------------------------

def authpatt(pa):
    f = p.read_csv(pa)
    f = f['Author(s) ID']
    authidlend = []
    for i in f.index: 
        t = str((f.iloc[i])).split(';')
        authidlend.append(len(t))
    authidlenu = list(set(authidlend))
    data = []
    for i in range(len(authidlenu)):
        data.append([])
        data[i].append(str(authidlenu[i]))
        data[i].append(authidlend.count(authidlenu[i]))
    df = p.DataFrame(data , columns = ['Number of Author affliation', 'Occurence'])             
    return (df)
#----------------------------------------------------------------------------------------------------
def prodpubl(pa):
    f = p.read_csv(pa)
    f = f['Publisher']
    f = f.dropna()
    l = list(f.drop_duplicates())
    ld = list(f)
    fin = []
    for i in range(len(l)):
        fin.append([])
        fin[i].append(l[i])
        fin[i].append(ld.count(l[i]))
    df = p.DataFrame(fin, columns = ['Publisher','Number of publications'] )
    df.sort_values(by = 'Number of publications', inplace = True, ascending = False)
    df = df.head(11)
    df.reset_index(drop = True, inplace = True)
    return (df)
#------------------------------------------------------------------------------------------------------    
#ffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffffff

w = Tk()
w.title('scientometrics ~ TRIGM Inc.')
w.geometry('6000x6000')

main_frame = Frame(w)
main_frame.pack(fill = BOTH , expand = 1)

canvas = Canvas(main_frame)
canvas.pack(side = LEFT, fill = BOTH, expand = 1) 

sc = Scrollbar(main_frame, orient = VERTICAL, command = canvas.yview)
sc.pack(side = RIGHT, fill = Y)

canvas.configure(yscrollcommand = sc.set)
canvas.bind('<Configure>', lambda e: canvas.configure(scrollregion = canvas.bbox("all")))

secondf = Frame(canvas,bg = 'black')
canvas.create_window((0,0), window = secondf, anchor = 'nw')

i = ''
def exit():
    w.destroy()   
def readfile():
    global f
    f = filedialog.askopenfilename(initialdir = 'home',filetypes = [('csv files','*.csv')])

but_browse =  Button(secondf,text = 'Browse',command = readfile,height = 1, width = 50 )
but_browse.pack(anchor = 'nw', padx = 465, pady = 3)

l = []
for i in range(7):
    l.append([])
    for j in range(3):
        l[i].append(False)
l[0][0] = IntVar()
l[0][1] = IntVar()
l[0][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Produtive year',font = 20, bg = 'grey')
frameprodyear.pack(anchor = W,pady = 3 , padx = 20)
c11 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[0][0], width = 128     )
c11.pack()
c12 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11,variable = l[0][1], width = 128     )
c12.pack()
c13 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[0][2], width = 128     )
c13.pack()

l[1][0] = IntVar()
l[1][1] = IntVar()
l[1][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Productive Authors',font = 20, bg = 'grey')
frameprodyear.pack(anchor = 'w',pady = 3 , padx = 20)
c21 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[1][0], width = 128     )
c21.pack()
c22 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[1][1], width = 128    )
c22.pack()
c23 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[1][2], width = 128    )
c23.pack()

l[2][0] = IntVar()
l[2][1] = IntVar()
l[2][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Produtive publisher',font = 20, bg = 'grey')
frameprodyear.pack(anchor = 'w',pady = 3 , padx = 20)
c31 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[2][0], width = 128     )
c31.pack()
c32 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[2][1], width = 128    )
c32.pack()
c33 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[2][2], width = 128    )
c33.pack()

l[3][0] = IntVar()
l[3][1] = IntVar()
l[3][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Authorship pattern',font = 20, bg = 'grey')
frameprodyear.pack(anchor = 'w',pady = 3 , padx = 20)
c41 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[3][0], width = 128     )
c41.pack()
c42 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[3][1], width = 128    )
c42.pack()
c43 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[3][2], width = 128    )
c43.pack()


l[4][0] = IntVar()
l[4][1] = IntVar()
l[4][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Author keywords',font = 20, bg = 'grey')
frameprodyear.pack(anchor = 'w',pady = 3 , padx = 20)
c51 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[4][0], width = 128    )
c51.pack()
c52 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[4][1], width = 128   )
c52.pack()
c53 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[4][2], width = 128     )
c53.pack()

l[5][0] = IntVar()
l[5][1] = IntVar()
l[5][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Index Kewords',font = 20, bg = 'grey')
frameprodyear.pack(anchor = 'w',pady = 3 , padx = 20)
c61 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[5][0], width = 128     )
c61.pack()
c62 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[5][1], width = 128    )
c62.pack()
c63 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[5][2], width = 128    )
c63.pack()


l[6][0] = IntVar()
l[6][1] = IntVar()
l[6][2] = IntVar()
frameprodyear = LabelFrame(secondf, text = 'Most Cited Authors',font = 20, bg = 'grey')
frameprodyear.pack(pady = 3)
c71 = Checkbutton(frameprodyear,anchor ='w', text = 'Bar',font = 11, variable = l[6][0], width = 128     )
c71.pack()
c72 = Checkbutton(frameprodyear,anchor ='w', text = 'Plot',font = 11, variable = l[6][1], width = 128    )
c72.pack()
c73 = Checkbutton(frameprodyear,anchor ='w', text = 'Pie',font = 11, variable = l[6][2], width = 128     )
c73.pack()

but_exit =  Button(secondf,text = 'exit',command = w.destroy,height = 1, width = 50 )
but_exit.pack(anchor = 'nw', padx = 465, pady = 3)

finlabel  = LabelFrame(secondf,bg = 'grey')
finlabel.pack(anchor = 'w', pady = 3 , padx = 20)
w.mainloop()
# 1 = bar 
# 2 = plot
# 3 = pie 

# c1 = prodyear
# c2 = prodauth
# c3 = prodauth
# c4 = prodpubl
# c5 = auhtkey
# c6 = indkey
# c7 = most cited authors
theobjs = []
for i in range(len(l)):
    for j in range(len(l[i])):
        if l[i][j].get() == 1:
            theobjs.append([i,j]) 
for i in range(len(theobjs)):
    if theobjs[i][0] == 0:
        if theobjs[i][1] == 0:    
            pbar(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pbar(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            pbar(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 1:
        if theobjs[i][1] == 0:    
            pbar(prodauth(f),f'temp{i}')
            todoc(prodauth(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pplolt(prodauth(f),f'temp{i}')
            todoc(prodauth(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(prodauth(f),f'temp{i}')
            todoc(prodauth(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 2:
        if theobjs[i][1] == 0:    
            pbar(prodpubl(f),f'temp{i}')
            todoc(prodpubl(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:        
            pplolt(prodpubl(f),f'temp{i}')
            todoc(prodpubl(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(prodpubl(f),f'temp{i}')
            todoc(prodpubl(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 3:
        if theobjs[i][1] == 0:    
            pbar(authpatt(f),f'temp{i}')
            todoc(authpatt(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pplolt(authpatt(f),f'temp{i}')
            todoc(authpatt(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(authpatt(f),f'temp{i}')
            todoc(authpatt(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 4:
        if theobjs[i][1] == 0:    
            pbar(authkeys(f),f'temp{i}')
            todoc(authkeys(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pplolt(authkeys(f),f'temp{i}')
            todoc(authkeys(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(authkeys(f),f'temp{i}')
            todoc(authkeys(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 5:
        if theobjs[i][1] == 0:    
            pbar(indkeys(f),f'temp{i}')
            todoc(indkeys(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pplolt(indkeys(f),f'temp{i}')
            todoc(indkeys(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(indkeys(f),f'temp{i}')
            todoc(indkeys(f),f'temp{i}','done.docx')
    elif theobjs[i][0] == 6:
        if theobjs[i][1] == 0:    
            pbar(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 1:    
            pplolt(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
        elif theobjs[i][1] == 2:    
            ppie(prodyear(f),f'temp{i}')
            todoc(prodyear(f),f'temp{i}','done.docx')
    
    