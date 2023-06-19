from docx import Document
from docx.shared import Inches
import pandas as p
d = Document()
def todoc(df,picturename,docname):    
    d.add_picture(f'{picturename}.png')
    d.add_page_break()
    l = []
    l.append(list(df.columns))
    for i in df.index:
        l.append(list(df.iloc[i]))     
    table = d.add_table( 0, (len(l[0])))
    for i,j in l:
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(j)
    d.add_page_break()
    d.save(docname)


